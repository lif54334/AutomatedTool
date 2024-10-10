'''
:@Author: lifei
:@Date: 2024/10/10 15:29:55
:@LastEditors: lifei
:@LastEditTime: 2024/10/10 15:29:55
:Description: 
:Copyright: Copyright (©) 2024 XXXX有限公司. All rights reserved.
'''
import datetime
import fitz  
import cv2
import os
import shutil
from docx import Document
from docxcompose.composer import Composer
from paddleocr import PPStructure,save_structure_res
from paddleocr.ppstructure.recovery.recovery_to_doc import sorted_layout_boxes, convert_info_docx
from copy import deepcopy
# 中文测试图
table_engine = PPStructure(recovery=True,lang='ch')

def add_page_break(doc):
    # 创建一个新的段落
    paragraph = doc.add_paragraph()
    # 为段落设置分页属性
    paragraph_format = paragraph.paragraph_format
    paragraph_format.page_break_before = True

def pdf2png(pdf_path, base_image_path):
    image_path=os.path.join(base_image_path,os.path.basename(pdf_path).split('.')[0])
    startTime_pdf2img = datetime.datetime.now()  # 开始时间
    print("imagePath=" + image_path)
    if not os.path.exists(image_path):
        os.makedirs(image_path)
    pdfDoc = fitz.open(pdf_path)
    totalPage=pdfDoc.page_count
    for pg in range(totalPage):
        page = pdfDoc[pg]
        rotate = int(0)
        zoom_x = 2
        zoom_y = 2
        mat = fitz.Matrix(zoom_x, zoom_y).prerotate(rotate)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        print(f'正在保存{pdf_path}的第{pg+1}页，共{totalPage}页')
        pix.save(image_path + '/' + f'images_{pg+1}.png')
    endTime_pdf2img = datetime.datetime.now()
    print(f'{pdfDoc}-pdf2img-花费时间={(endTime_pdf2img - startTime_pdf2img).seconds}秒')

def img2docx(img_path,word_path):
    master = Document()
    composer = Composer(master)
    imgs=os.listdir(img_path)
    for img_name in imgs:
        print(os.path.join(img_path,img_name))
        img = cv2.imread(os.path.join(img_path,img_name))
        result = table_engine(img)

        save_structure_res(result, save_folder, os.path.basename(img_path).split('.')[0])

        h, w, _ = img.shape
        res = sorted_layout_boxes(result, w)
        convert_info_docx(img, res, save_folder, os.path.basename(img_path).split('.')[0])
        doc = Document('./output/word/tmp_ocr.docx')
        # add_page_break(doc)
        composer.append(doc)
    composer.save(word_path)

def find_pdf_files(directory):
    pdf_files = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('.pdf'):
                pdf_files.append(os.path.join(root, file))
    return pdf_files

def clear_directories():
    # 定义需要清空的目录路径
    directories = ["./output/imgs/tmp", "./output/word/tmp"]
    
    # 遍历每个目录
    for directory in directories:
        # 检查目录是否存在
        if os.path.exists(directory):
            try:
                # 使用shutil.rmtree递归删除目录下的所有内容
                shutil.rmtree(directory)
                print(f"Directory {directory} has been cleared.")
            except OSError as e:
                print(f"Error: {e.strerror}")
        else:
            print(f"Directory {directory} does not exist.")


def move_and_rename_pdf(src_path, new_name, target_dir):
    """
    重命名并移动PDF文件到指定目录。

    参数:
    src_path (str): 源PDF文件的完整路径。
    new_name (str): 目标PDF文件的新名称（不含路径）。
    target_dir (str): 目标目录的路径。

    返回:
    bool: 操作成功返回True，否则返回False。
    """
    try:
        # 确保目标目录存在
        if not os.path.exists(target_dir):
            os.makedirs(target_dir)
        
        # 构建目标文件的完整路径
        new_file_path = os.path.join(target_dir, new_name)
        
        # 使用shutil.move进行文件的移动和重命名
        shutil.copy(src_path, new_file_path)
        
        print(f"文件已成功移动并重命名为: {new_file_path}")
    except FileNotFoundError:
        print("原始文件未找到，请检查路径是否正确。")
    except PermissionError:
        print("没有足够的权限访问文件或目录。")
    except Exception as e:
        print(f"移动文件时发生未知错误: {e}")


if __name__ == "__main__":
    save_folder = './output/word'
    pdfPath = './input/pdf'
    base_image_path = './output/imgs'
    image_path = './output/imgs/tmp'
    
    pdf_files = find_pdf_files(pdfPath)
    for file in pdf_files:
        clear_directories()
        print(f"文件名: {os.path.basename(file)}, 文件路径: {file}")
        move_and_rename_pdf(file, 'tmp.pdf', './input/tmp_pdf')
        new_file_path = './input/tmp_pdf/tmp.pdf'

        pdf2png(new_file_path, base_image_path)
        word_path='./output/word/result/{}.docx'.format(os.path.basename(file).split('.')[0])
        img2docx(image_path,word_path)
        # clear_directories()